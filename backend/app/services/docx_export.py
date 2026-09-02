# -*- coding: utf-8 -*-
"""把修订版合同 + 修改说明清单导出成 .docx（python-docx）。
结构：修订版合同全文 → 修改说明表（原条款/修订后/依据）→ 引用校验脚注。
"""
import io

from docx import Document


def build_docx(revised: str, changes: list, ok: int, total: int) -> bytes:
    doc = Document()
    doc.add_heading("修订版合同", level=0)
    for para in revised.splitlines():
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.add_heading("修改说明", level=1)
    if changes:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ("原条款", "修订后", "依据")):
            cell.text = text
        for c in changes:
            row = table.add_row().cells
            row[0].text = c.get("原条款", "")
            row[1].text = c.get("修订后", "")
            mark = "✅" if c.get("依据真实") else "⚠️ 未核实"
            row[2].text = f"{c.get('依据', '')} {mark}"
    else:
        doc.add_paragraph("（无修改项）")

    if total == 0:
        note = "引用校验：无修改项，无需核对依据。"
    elif ok == total:
        note = f"引用校验：{total} 处修改依据均真实存在（来源：检索语料）"
    else:
        note = f"引用校验：{total - ok} 处修改依据未能在语料核实，已如实标注 ⚠️。"
    doc.add_paragraph("")
    doc.add_paragraph(note)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
