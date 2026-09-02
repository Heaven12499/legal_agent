# -*- coding: utf-8 -*-
"""
文件 → 纯文本：支持合同审查的上传输入（.docx / .pdf / .txt）。

只做"提取文本"一件事，不做任何审查逻辑——提取结果由前端填回输入框，
再走现有 /api/chat 复用同一个 agent。红线：只读不写，不执行宏，不信任内容。
"""
from pathlib import Path

# 允许的扩展名（小写）→ 提取函数
SUPPORTED = {".docx", ".pdf", ".txt", ".md"}

MAX_CHARS = 200_000  # 上限保护：超长文档截断，避免上下文爆炸


def extract_text(data: bytes, filename: str) -> str:
    """按扩展名分发提取纯文本。未知/不支持格式抛 ValueError。"""
    ext = Path(filename).suffix.lower()
    if ext == ".docx":
        text = _from_docx(data)
    elif ext == ".pdf":
        text = _from_pdf(data)
    elif ext in (".txt", ".md"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"不支持的文件格式：{ext or '(无扩展名)'}，请上传 .docx / .pdf / .txt")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    return text[:MAX_CHARS]


def _from_docx(data: bytes) -> str:
    """docx = zip 容器，python-docx 只读解包；宏（vba）不执行，天然安全。"""
    import io

    from docx import Document

    doc = Document(io.BytesIO(data))
    # 段落 + 表格都要，合同常把条款放表格里
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_pdf(data: bytes) -> str:
    """纯 python 读文本层；扫描版（无文本层）返回空，前端提示改上传 docx/txt。"""
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)
